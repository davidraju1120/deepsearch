import os
import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
import json
import re
from datetime import datetime
from dataclasses import dataclass

import docx
import PyPDF2
import markdown
from bs4 import BeautifulSoup
import requests
from tqdm import tqdm

@dataclass
class ProcessedDocument:
    """Represents a processed document with metadata."""
    content: str
    metadata: Dict[str, Any]
    source_path: str
    processing_timestamp: datetime
    word_count: int
    char_count: int

class DocumentProcessor:
    """
    Document processor for ingesting various file formats and web content.
    """
    
    def __init__(self, supported_formats: Optional[List[str]] = None):
        """
        Initialize the document processor.
        
        Args:
            supported_formats: List of supported file formats
        """
        self.supported_formats = supported_formats or [
            '.txt', '.md', '.pdf', '.docx', '.html', '.json'
        ]
        
        self.processing_stats = {
            'total_processed': 0,
            'by_format': {},
            'errors': 0
        }
        
        logging.info(f"DocumentProcessor initialized with formats: {self.supported_formats}")
    
    def process_file(self, file_path: str, 
                    metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process a single file.
        
        Args:
            file_path: Path to the file to process
            metadata: Additional metadata to include
            
        Returns:
            ProcessedDocument object
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Initialize metadata
        if metadata is None:
            metadata = {}
        
        # Add file metadata
        metadata.update({
            'source_file': str(file_path),
            'file_extension': file_extension,
            'file_size': file_path.stat().st_size,
            'processed_at': datetime.now().isoformat()
        })
        
        try:
            # Process based on file type
            if file_extension == '.txt':
                content = self._process_txt_file(file_path)
            elif file_extension == '.md':
                content = self._process_markdown_file(file_path)
            elif file_extension == '.pdf':
                content = self._process_pdf_file(file_path)
            elif file_extension == '.docx':
                content = self._process_docx_file(file_path)
            elif file_extension == '.html':
                content = self._process_html_file(file_path)
            elif file_extension == '.json':
                content = self._process_json_file(file_path)
            else:
                raise ValueError(f"Handler not implemented for {file_extension}")
            
            # Create processed document
            processed_doc = ProcessedDocument(
                content=content,
                metadata=metadata,
                source_path=str(file_path),
                processing_timestamp=datetime.now(),
                word_count=len(content.split()),
                char_count=len(content)
            )
            
            # Update statistics
            self.processing_stats['total_processed'] += 1
            self.processing_stats['by_format'][file_extension] = \
                self.processing_stats['by_format'].get(file_extension, 0) + 1
            
            logging.info(f"Successfully processed file: {file_path}")
            return processed_doc
            
        except Exception as e:
            self.processing_stats['errors'] += 1
            logging.error(f"Error processing file {file_path}: {e}")
            raise
    
    def process_directory(self, directory_path: str, 
                         recursive: bool = True,
                         metadata: Optional[Dict[str, Any]] = None) -> List[ProcessedDocument]:
        """
        Process all files in a directory.
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to process subdirectories
            metadata: Additional metadata to include
            
        Returns:
            List of ProcessedDocument objects
        """
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        processed_docs = []
        
        # Get all files
        if recursive:
            files = [f for f in directory_path.rglob('*') if f.is_file()]
        else:
            files = [f for f in directory_path.iterdir() if f.is_file()]
        
        # Filter by supported formats
        supported_files = [f for f in files if f.suffix.lower() in self.supported_formats]
        
        logging.info(f"Found {len(supported_files)} supported files in {directory_path}")
        
        # Process files with progress bar
        for file_path in tqdm(supported_files, desc="Processing files"):
            try:
                processed_doc = self.process_file(str(file_path), metadata)
                processed_docs.append(processed_doc)
            except Exception as e:
                logging.warning(f"Failed to process {file_path}: {e}")
                continue
        
        logging.info(f"Successfully processed {len(processed_docs)} files from {directory_path}")
        return processed_docs
    
    def process_url(self, url: str, 
                   metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process content from a URL.
        
        Args:
            url: URL to process
            metadata: Additional metadata to include
            
        Returns:
            ProcessedDocument object
        """
        if metadata is None:
            metadata = {}
        
        metadata.update({
            'source_url': url,
            'processed_at': datetime.now().isoformat()
        })
        
        try:
            # Fetch content
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract text content
            content = self._extract_text_from_html(soup)
            
            # Add URL metadata
            metadata.update({
                'content_type': response.headers.get('content-type', ''),
                'status_code': response.status_code,
                'final_url': response.url
            })
            
            processed_doc = ProcessedDocument(
                content=content,
                metadata=metadata,
                source_path=url,
                processing_timestamp=datetime.now(),
                word_count=len(content.split()),
                char_count=len(content)
            )
            
            self.processing_stats['total_processed'] += 1
            self.processing_stats['by_format']['url'] = \
                self.processing_stats['by_format'].get('url', 0) + 1
            
            logging.info(f"Successfully processed URL: {url}")
            return processed_doc
            
        except Exception as e:
            self.processing_stats['errors'] += 1
            logging.error(f"Error processing URL {url}: {e}")
            raise
    
    def process_text(self, text: str, 
                    metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process raw text content.
        
        Args:
            text: Text content to process
            metadata: Additional metadata to include
            
        Returns:
            ProcessedDocument object
        """
        if metadata is None:
            metadata = {}
        
        metadata.update({
            'source_type': 'text',
            'processed_at': datetime.now().isoformat()
        })
        
        # Clean and preprocess text
        cleaned_text = self._clean_text(text)
        
        processed_doc = ProcessedDocument(
            content=cleaned_text,
            metadata=metadata,
            source_path='text_input',
            processing_timestamp=datetime.now(),
            word_count=len(cleaned_text.split()),
            char_count=len(cleaned_text)
        )
        
        self.processing_stats['total_processed'] += 1
        self.processing_stats['by_format']['text'] = \
            self.processing_stats['by_format'].get('text', 0) + 1
        
        return processed_doc
    
    # File format specific processors
    def _process_txt_file(self, file_path: Path) -> str:
        """Process a text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return self._clean_text(content)
    
    def _process_markdown_file(self, file_path: Path) -> str:
        """Process a markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Convert markdown to plain text
        html_content = markdown.markdown(md_content)
        soup = BeautifulSoup(html_content, 'html.parser')
        return self._extract_text_from_html(soup)
    
    def _process_pdf_file(self, file_path: Path) -> str:
        """Process a PDF file."""
        content = []
        
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    content.append(page_text)
        
        return self._clean_text('\n'.join(content))
    
    def _process_docx_file(self, file_path: Path) -> str:
        """Process a DOCX file."""
        doc = docx.Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        return self._clean_text('\n'.join(content))
    
    def _process_html_file(self, file_path: Path) -> str:
        """Process an HTML file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        return self._extract_text_from_html(soup)
    
    def _process_json_file(self, file_path: Path) -> str:
        """Process a JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
        
        # Convert JSON to text representation
        if isinstance(json_data, dict):
            content = json.dumps(json_data, indent=2)
        elif isinstance(json_data, list):
            content = '\n'.join(json.dumps(item, indent=2) for item in json_data)
        else:
            content = str(json_data)
        
        return self._clean_text(content)
    
    # Helper methods
    def _extract_text_from_html(self, soup: BeautifulSoup) -> str:
        """Extract clean text from HTML soup."""
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean and return
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove empty lines
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Join with newlines
        cleaned_text = '\n'.join(lines)
        
        # Remove special characters that might cause issues
        cleaned_text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\^\&\*\+\=\~\`]', '', cleaned_text)
        
        return cleaned_text.strip()
    
    def chunk_document(self, document: ProcessedDocument, 
                      chunk_size: int = 1000, 
                      overlap: int = 200) -> List[ProcessedDocument]:
        """
        Split a document into smaller chunks.
        
        Args:
            document: Document to chunk
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of chunked ProcessedDocument objects
        """
        content = document.content
        words = content.split()
        
        if len(words) <= chunk_size:
            return [document]
        
        chunks = []
        start_idx = 0
        
        while start_idx < len(words):
            end_idx = min(start_idx + chunk_size, len(words))
            chunk_words = words[start_idx:end_idx]
            chunk_content = ' '.join(chunk_words)
            
            # Create chunk metadata
            chunk_metadata = document.metadata.copy()
            chunk_metadata.update({
                'chunk_id': len(chunks),
                'chunk_start': start_idx,
                'chunk_end': end_idx,
                'is_chunk': True,
                'original_document': document.source_path
            })
            
            chunk_doc = ProcessedDocument(
                content=chunk_content,
                metadata=chunk_metadata,
                source_path=f"{document.source_path}_chunk_{len(chunks)}",
                processing_timestamp=datetime.now(),
                word_count=len(chunk_words),
                char_count=len(chunk_content)
            )
            
            chunks.append(chunk_doc)
            start_idx = end_idx - overlap
        
        return chunks
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics."""
        return self.processing_stats.copy()
    
    def reset_statistics(self):
        """Reset processing statistics."""
        self.processing_stats = {
            'total_processed': 0,
            'by_format': {},
            'errors': 0
        }
        logging.info("Processing statistics reset")

class DocumentIngestor:
    """
    High-level document ingestion system that combines processing and storage.
    """
    
    def __init__(self, document_store, processor: Optional[DocumentProcessor] = None):
        """
        Initialize the document ingestor.
        
        Args:
            document_store: DocumentStore instance
            processor: DocumentProcessor instance
        """
        self.document_store = document_store
        self.processor = processor or DocumentProcessor()
        
        self.ingestion_stats = {
            'total_ingested': 0,
            'by_source': {},
            'errors': 0
        }
    
    def ingest_file(self, file_path: str, 
                   metadata: Optional[Dict[str, Any]] = None,
                   chunk_document: bool = False,
                   chunk_size: int = 1000,
                   chunk_overlap: int = 200) -> List[str]:
        """
        Ingest a single file into the document store.
        
        Args:
            file_path: Path to the file
            metadata: Additional metadata
            chunk_document: Whether to chunk the document
            chunk_size: Size of chunks if chunking
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of document IDs
        """
        try:
            # Process the file
            processed_doc = self.processor.process_file(file_path, metadata)
            
            # Chunk if requested
            if chunk_document:
                chunks = self.processor.chunk_document(processed_doc, chunk_size, chunk_overlap)
                doc_ids = []
                for chunk in chunks:
                    doc_id = self.document_store.add_document(
                        chunk.content, 
                        chunk.metadata, 
                        chunk.source_path
                    )
                    doc_ids.append(doc_id)
            else:
                doc_id = self.document_store.add_document(
                    processed_doc.content, 
                    processed_doc.metadata, 
                    processed_doc.source_path
                )
                doc_ids = [doc_id]
            
            # Update statistics
            self.ingestion_stats['total_ingested'] += len(doc_ids)
            self.ingestion_stats['by_source']['file'] = \
                self.ingestion_stats['by_source'].get('file', 0) + len(doc_ids)
            
            logging.info(f"Successfully ingested file {file_path} with {len(doc_ids)} documents")
            return doc_ids
            
        except Exception as e:
            self.ingestion_stats['errors'] += 1
            logging.error(f"Error ingesting file {file_path}: {e}")
            raise
    
    def ingest_directory(self, directory_path: str,
                        recursive: bool = True,
                        metadata: Optional[Dict[str, Any]] = None,
                        chunk_documents: bool = False,
                        chunk_size: int = 1000,
                        chunk_overlap: int = 200) -> List[str]:
        """
        Ingest all files in a directory.
        
        Args:
            directory_path: Path to the directory
            recursive: Whether to process subdirectories
            metadata: Additional metadata
            chunk_documents: Whether to chunk documents
            chunk_size: Size of chunks if chunking
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of all document IDs
        """
        try:
            # Process all files
            processed_docs = self.processor.process_directory(directory_path, recursive, metadata)
            
            all_doc_ids = []
            
            # Ingest each document
            for processed_doc in processed_docs:
                if chunk_documents:
                    chunks = self.processor.chunk_document(processed_doc, chunk_size, chunk_overlap)
                    for chunk in chunks:
                        doc_id = self.document_store.add_document(
                            chunk.content, 
                            chunk.metadata, 
                            chunk.source_path
                        )
                        all_doc_ids.append(doc_id)
                else:
                    doc_id = self.document_store.add_document(
                        processed_doc.content, 
                        processed_doc.metadata, 
                        processed_doc.source_path
                    )
                    all_doc_ids.append(doc_id)
            
            # Update statistics
            self.ingestion_stats['total_ingested'] += len(all_doc_ids)
            self.ingestion_stats['by_source']['directory'] = \
                self.ingestion_stats['by_source'].get('directory', 0) + len(all_doc_ids)
            
            logging.info(f"Successfully ingested {len(all_doc_ids)} documents from {directory_path}")
            return all_doc_ids
            
        except Exception as e:
            self.ingestion_stats['errors'] += 1
            logging.error(f"Error ingesting directory {directory_path}: {e}")
            raise
    
    def ingest_text(self, text: str, 
                   metadata: Optional[Dict[str, Any]] = None,
                   doc_id: Optional[str] = None) -> str:
        """
        Ingest raw text content.
        
        Args:
            text: Text content to ingest
            metadata: Additional metadata
            doc_id: Optional document ID
            
        Returns:
            Document ID
        """
        try:
            processed_doc = self.processor.process_text(text, metadata)
            
            doc_id = self.document_store.add_document(
                processed_doc.content, 
                processed_doc.metadata, 
                doc_id
            )
            
            # Update statistics
            self.ingestion_stats['total_ingested'] += 1
            self.ingestion_stats['by_source']['text'] = \
                self.ingestion_stats['by_source'].get('text', 0) + 1
            
            logging.info(f"Successfully ingested text document with ID: {doc_id}")
            return doc_id
            
        except Exception as e:
            self.ingestion_stats['errors'] += 1
            logging.error(f"Error ingesting text: {e}")
            raise
    
    def ingest_url(self, url: str, 
                  metadata: Optional[Dict[str, Any]] = None,
                  chunk_document: bool = False,
                  chunk_size: int = 1000,
                  chunk_overlap: int = 200) -> List[str]:
        """
        Ingest content from a URL.
        
        Args:
            url: URL to ingest
            metadata: Additional metadata
            chunk_document: Whether to chunk the document
            chunk_size: Size of chunks if chunking
            chunk_overlap: Overlap between chunks
            
        Returns:
            List of document IDs
        """
        try:
            processed_doc = self.processor.process_url(url, metadata)
            
            if chunk_document:
                chunks = self.processor.chunk_document(processed_doc, chunk_size, chunk_overlap)
                doc_ids = []
                for chunk in chunks:
                    doc_id = self.document_store.add_document(
                        chunk.content, 
                        chunk.metadata, 
                        chunk.source_path
                    )
                    doc_ids.append(doc_id)
            else:
                doc_id = self.document_store.add_document(
                    processed_doc.content, 
                    processed_doc.metadata, 
                    processed_doc.source_path
                )
                doc_ids = [doc_id]
            
            # Update statistics
            self.ingestion_stats['total_ingested'] += len(doc_ids)
            self.ingestion_stats['by_source']['url'] = \
                self.ingestion_stats['by_source'].get('url', 0) + len(doc_ids)
            
            logging.info(f"Successfully ingested URL {url} with {len(doc_ids)} documents")
            return doc_ids
            
        except Exception as e:
            self.ingestion_stats['errors'] += 1
            logging.error(f"Error ingesting URL {url}: {e}")
            raise
    
    def get_ingestion_statistics(self) -> Dict[str, Any]:
        """Get ingestion statistics."""
        stats = self.ingestion_stats.copy()
        stats['processing_stats'] = self.processor.get_processing_statistics()
        stats['document_store_stats'] = self.document_store.get_statistics()
        return stats
    
    def reset_statistics(self):
        """Reset ingestion statistics."""
        self.ingestion_stats = {
            'total_ingested': 0,
            'by_source': {},
            'errors': 0
        }
        self.processor.reset_statistics()
        logging.info("Ingestion statistics reset")
